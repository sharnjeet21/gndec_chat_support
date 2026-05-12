"""
GNDEC Deep Scraper v2
======================
- Crawls gndec.ac.in + all sub-domains up to MAX_DEPTH levels deep
- Downloads and parses PDF, DOCX, DOC files into text
- Converts all content into Q&A pairs for the RAG vector DB

Output: data/gndec_data.json

Usage:
    python3 scraper/gndec_scraper.py
"""

import io
import json
import os
import re
import time
import logging
import hashlib
import tempfile
from collections import deque
from urllib.parse import urljoin, urlparse
from typing import List, Dict, Any, Set, Tuple, Optional

import requests
from bs4 import BeautifulSoup

# ── PDF / DOCX parsers ────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ─────────────────────────────────────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────
ROOT_URL             = "https://gndec.ac.in"
ALLOWED_DOMAIN_SUFFIX = "gndec.ac.in"
MAX_DEPTH            = 2          # crawl depth per site
MAX_PAGES_PER_SITE   = 120        # safety cap per subdomain
MAX_DOCS_PER_SITE    = 30         # max PDF/DOCX files per subdomain
REQUEST_TIMEOUT      = 20         # seconds
DELAY_HTML           = 0.6        # seconds between HTML requests
DELAY_DOC            = 1.2        # seconds between document downloads
MAX_DOC_SIZE_MB      = 10         # skip files larger than this
CHUNK_SIZE           = 800        # chars per text chunk for long docs

SEED_SITES = [
    "https://gndec.ac.in",
    "https://ee.gndec.ac.in",
    "https://it.gndec.ac.in",
    "https://cse.gndec.ac.in",
    "https://ce.gndec.ac.in",
    "https://me.gndec.ac.in",
    "https://ece.gndec.ac.in",
    "https://mca.gndec.ac.in",
    "https://mba.gndec.ac.in",
    "https://admission.gndec.ac.in",
    "https://academics.gndec.ac.in",
    "https://tcc.gndec.ac.in",
]

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}

# Extensions to download and parse as documents
DOC_EXTENSIONS = {".pdf", ".doc", ".docx"}

# Extensions to completely skip (images, archives, media)
SKIP_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".gif", ".svg", ".ico", ".webp", ".bmp",
    ".mp4", ".mp3", ".avi", ".mov", ".wmv",
    ".zip", ".rar", ".tar", ".gz", ".7z",
    ".exe", ".dmg", ".apk",
    ".xls", ".xlsx", ".ppt", ".pptx",  # skip spreadsheets/slides (hard to parse meaningfully)
    ".css", ".js", ".json", ".xml",
}

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def url_extension(url: str) -> str:
    path = urlparse(url).path.lower()
    _, ext = os.path.splitext(path)
    return ext


def is_doc_url(url: str) -> bool:
    return url_extension(url) in DOC_EXTENSIONS


def is_html_url(url: str) -> bool:
    ext = url_extension(url)
    return ext == "" or ext in (".html", ".htm", ".php", ".asp", ".aspx")


def is_skippable(url: str) -> bool:
    return url_extension(url) in SKIP_EXTENSIONS


def is_gndec_url(url: str) -> bool:
    try:
        p = urlparse(url)
        return (
            p.scheme in ("http", "https")
            and p.netloc.endswith(ALLOWED_DOMAIN_SUFFIX)
        )
    except Exception:
        return False


def normalise_url(url: str) -> str:
    return url.split("#")[0].rstrip("/")


def get_origin(url: str) -> str:
    p = urlparse(url)
    return f"{p.scheme}://{p.netloc}"


def url_hash(url: str) -> str:
    return hashlib.md5(url.lower().encode()).hexdigest()


def text_hash(text: str) -> str:
    return hashlib.md5(text.lower().strip().encode()).hexdigest()


# ─────────────────────────────────────────────────────────────────────────────
# HTTP FETCH
# ─────────────────────────────────────────────────────────────────────────────

def fetch_html(url: str) -> Tuple[Optional[str], Optional[BeautifulSoup]]:
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT,
                         allow_redirects=True)
        r.raise_for_status()
        ct = r.headers.get("content-type", "")
        if "html" not in ct and "text" not in ct:
            return None, None
        soup = BeautifulSoup(r.text, "html.parser")
        return r.url, soup
    except Exception as e:
        logger.debug(f"HTML fetch failed {url}: {e}")
        return None, None


def fetch_bytes(url: str) -> Optional[bytes]:
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT,
                         allow_redirects=True, stream=True)
        r.raise_for_status()
        # Check size
        cl = r.headers.get("content-length")
        if cl and int(cl) > MAX_DOC_SIZE_MB * 1024 * 1024:
            logger.info(f"  Skipping large file ({int(cl)//1024//1024}MB): {url}")
            return None
        data = b""
        for chunk in r.iter_content(65536):
            data += chunk
            if len(data) > MAX_DOC_SIZE_MB * 1024 * 1024:
                logger.info(f"  Skipping oversized file: {url}")
                return None
        return data
    except Exception as e:
        logger.debug(f"Doc fetch failed {url}: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT PARSERS
# ─────────────────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes, url: str) -> str:
    """Extract text from PDF bytes. Tries PyMuPDF first, falls back to pdfplumber."""
    text = ""

    if HAS_PYMUPDF:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            pages = []
            for page in doc:
                pages.append(page.get_text("text"))
            text = "\n\n".join(pages)
            doc.close()
            if text.strip():
                logger.info(f"  PDF parsed via PyMuPDF: {len(text)} chars")
                return clean_text(text)
        except Exception as e:
            logger.debug(f"PyMuPDF failed for {url}: {e}")

    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            text = "\n\n".join(pages)
            if text.strip():
                logger.info(f"  PDF parsed via pdfplumber: {len(text)} chars")
                return clean_text(text)
        except Exception as e:
            logger.debug(f"pdfplumber failed for {url}: {e}")

    logger.warning(f"  Could not parse PDF: {url}")
    return ""


def parse_docx(data: bytes, url: str) -> str:
    """Extract text from DOCX bytes."""
    if not HAS_DOCX:
        logger.warning("python-docx not installed, skipping DOCX")
        return ""
    try:
        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n\n".join(paragraphs)
        logger.info(f"  DOCX parsed: {len(text)} chars")
        return clean_text(text)
    except Exception as e:
        logger.debug(f"DOCX parse failed for {url}: {e}")
        return ""


def parse_doc_legacy(data: bytes, url: str) -> str:
    """Try to extract text from old .doc files using antiword or catdoc if available."""
    try:
        import subprocess
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(data)
            tmp_path = f.name
        result = subprocess.run(
            ["antiword", tmp_path], capture_output=True, text=True, timeout=15
        )
        os.unlink(tmp_path)
        if result.returncode == 0 and result.stdout.strip():
            return clean_text(result.stdout)
    except Exception:
        pass
    return ""


def parse_document(data: bytes, url: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = url_extension(url)
    if ext == ".pdf":
        return parse_pdf(data, url)
    elif ext == ".docx":
        return parse_docx(data, url)
    elif ext == ".doc":
        return parse_doc_legacy(data, url)
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# TEXT → Q&A CHUNKS
# ─────────────────────────────────────────────────────────────────────────────

def chunk_text_to_qa(text: str, source_url: str, doc_title: str,
                     section: str) -> List[Dict[str, Any]]:
    """
    Split a long document text into overlapping chunks and
    wrap each as a Q&A pair so the RAG system can retrieve it.
    """
    source_label = urlparse(source_url).netloc or source_url
    pairs = []
    sentences = re.split(r"(?<=[.!?])\s+|\n{2,}", text)

    current_chunk = []
    current_len   = 0

    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue
        current_chunk.append(sent)
        current_len += len(sent)

        if current_len >= CHUNK_SIZE:
            chunk_text = " ".join(current_chunk)
            # Use first sentence as the "question" context
            first_sent = current_chunk[0][:120]
            q = f"What does the document say about: {first_sent}?"
            pairs.append({
                "question": q,
                "answer": chunk_text[:2000],
                "section": section or doc_title,
                "source_file": source_label,
                "doc_url": source_url,
            })
            # Overlap: keep last 2 sentences for context continuity
            current_chunk = current_chunk[-2:]
            current_len   = sum(len(s) for s in current_chunk)

    # Flush remaining
    if current_chunk and current_len > 50:
        chunk_text = " ".join(current_chunk)
        first_sent = current_chunk[0][:120]
        q = f"What does the document say about: {first_sent}?"
        pairs.append({
            "question": q,
            "answer": chunk_text[:2000],
            "section": section or doc_title,
            "source_file": source_label,
            "doc_url": source_url,
        })

    return pairs


def doc_text_to_qa(text: str, url: str, page_title: str) -> List[Dict[str, Any]]:
    """
    Convert parsed document text into Q&A pairs.
    Strategy:
      1. If text has clear headings → heading + body as Q&A
      2. Otherwise → chunk into fixed-size overlapping windows
    """
    if not text or len(text) < 50:
        return []

    source_label = urlparse(url).netloc or url
    pairs: List[Dict[str, Any]] = []

    # Try heading-based splitting first
    heading_pattern = re.compile(
        r"^(?:\d+[\.\)]\s+|[A-Z][A-Z\s]{3,}:?\s*$)", re.MULTILINE
    )
    sections = heading_pattern.split(text)
    headings = heading_pattern.findall(text)

    if len(headings) >= 2:
        for heading, body in zip(headings, sections[1:]):
            heading = clean_text(heading)
            body    = clean_text(body)
            if not heading or len(body) < 30:
                continue
            # Chunk body if too long
            if len(body) <= 1500:
                q = f"What is {heading}?" if not heading.endswith("?") else heading
                pairs.append({
                    "question": q,
                    "answer": body,
                    "section": page_title,
                    "source_file": source_label,
                    "doc_url": url,
                })
            else:
                pairs.extend(chunk_text_to_qa(body, url, heading, page_title))
    else:
        # No clear headings — chunk the whole text
        pairs.extend(chunk_text_to_qa(text, url, page_title, page_title))

    return pairs


# ─────────────────────────────────────────────────────────────────────────────
# HTML CONTENT EXTRACTION → Q&A
# ─────────────────────────────────────────────────────────────────────────────

def extract_html_qa(soup: BeautifulSoup, url: str) -> List[Dict[str, Any]]:
    """Extract Q&A pairs from an HTML page."""
    pairs: List[Dict[str, Any]] = []

    for tag in soup(["script", "style", "nav", "footer", "noscript", "aside", "iframe"]):
        tag.decompose()

    page_title = clean_text(soup.title.string) if (soup.title and soup.title.string) else ""
    source_label = urlparse(url).netloc

    if not page_title:
        h1 = soup.find("h1")
        page_title = clean_text(h1.get_text()) if h1 else source_label

    # ── 1. Definition lists ───────────────────────────────────────────────
    for dl in soup.find_all("dl"):
        for dt, dd in zip(dl.find_all("dt"), dl.find_all("dd")):
            q = clean_text(dt.get_text())
            a = clean_text(dd.get_text())
            if q and a and len(a) > 10:
                pairs.append({"question": q, "answer": a,
                               "section": page_title, "source_file": source_label})

    # ── 2. Tables ─────────────────────────────────────────────────────────
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue
        # Get headers from first row
        header_cells = rows[0].find_all(["th", "td"])
        headers = [clean_text(c.get_text()) for c in header_cells]

        for row in rows[1:]:
            cells = row.find_all(["td", "th"])
            if not cells:
                continue
            vals = [clean_text(c.get_text()) for c in cells]

            if len(vals) == 2 and vals[0] and vals[1]:
                pairs.append({"question": vals[0], "answer": vals[1],
                               "section": page_title, "source_file": source_label})
            elif len(vals) >= 2 and vals[0]:
                a = " | ".join(
                    f"{h}: {v}" for h, v in zip(headers, vals) if h and v
                )
                if len(a) > 15:
                    pairs.append({
                        "question": f"What are the details for: {vals[0]}?",
                        "answer": a,
                        "section": page_title,
                        "source_file": source_label,
                    })

    # ── 3. Headings + body paragraphs ─────────────────────────────────────
    main = (
        soup.find("main")
        or soup.find(id=re.compile(r"content|main|body", re.I))
        or soup.find(class_=re.compile(r"content|main|body|wrapper|entry", re.I))
        or soup.body
    )

    if main:
        for heading in main.find_all(["h1", "h2", "h3", "h4"]):
            h_text = clean_text(heading.get_text())
            if not h_text or len(h_text) < 4:
                continue

            body_parts = []
            for sib in heading.find_next_siblings():
                if sib.name in ("h1", "h2", "h3", "h4"):
                    break
                t = clean_text(sib.get_text())
                if t and len(t) > 20:
                    body_parts.append(t)
                if len(body_parts) >= 6:
                    break

            if body_parts:
                answer = " ".join(body_parts)
                if len(answer) > 2000:
                    answer = answer[:2000] + "…"
                q = f"What is {h_text}?" if not h_text.endswith("?") else h_text
                pairs.append({"question": q, "answer": answer,
                               "section": page_title, "source_file": source_label})

    # ── 4. Fallback: meta description + first paragraphs ─────────────────
    if not pairs:
        meta = soup.find("meta", attrs={"name": re.compile(r"description", re.I)})
        meta_text = clean_text(meta.get("content", "")) if meta else ""
        body_text = " ".join(
            clean_text(p.get_text())
            for p in (soup.find_all("p") or [])[:6]
            if len(clean_text(p.get_text())) > 30
        )
        answer = meta_text or body_text
        if answer and len(answer) > 20:
            pairs.append({
                "question": f"What is {page_title}?",
                "answer": answer[:2000],
                "section": page_title,
                "source_file": source_label,
            })

    return pairs


# ─────────────────────────────────────────────────────────────────────────────
# LINK EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_links(soup: BeautifulSoup, base_url: str, origin: str
                  ) -> Tuple[Set[str], Set[str]]:
    """
    Returns (html_links, doc_links) — both filtered to same origin.
    """
    html_links: Set[str] = set()
    doc_links:  Set[str] = set()

    for tag in soup.find_all("a", href=True):
        href = tag["href"].strip()
        if not href or href.startswith("javascript") or href.startswith("mailto"):
            continue
        full = normalise_url(urljoin(base_url, href))
        if not is_gndec_url(full):
            continue
        if get_origin(full) != origin:
            continue

        if is_skippable(full):
            continue
        elif is_doc_url(full):
            doc_links.add(full)
        elif is_html_url(full):
            html_links.add(full)

    return html_links, doc_links


def discover_subdomains(soup: BeautifulSoup, base_url: str,
                        known: Set[str]) -> Set[str]:
    """Find new gndec.ac.in subdomains linked from a page."""
    found: Set[str] = set()
    base_origin = get_origin(base_url)
    for tag in soup.find_all("a", href=True):
        href = tag["href"].strip()
        full = normalise_url(urljoin(base_url, href))
        if not is_gndec_url(full):
            continue
        o = get_origin(full)
        if o != base_origin and o not in known:
            found.add(o)
    return found


# ─────────────────────────────────────────────────────────────────────────────
# SITE CRAWLER  (BFS up to MAX_DEPTH)
# ─────────────────────────────────────────────────────────────────────────────

def crawl_site(root: str,
               visited_urls: Set[str],
               visited_docs: Set[str]) -> List[Dict[str, Any]]:
    """
    BFS crawl of a single subdomain up to MAX_DEPTH.
    Collects HTML Q&A pairs AND parses linked documents.
    """
    origin     = get_origin(root)
    site_label = urlparse(root).netloc
    all_pairs: List[Dict[str, Any]] = []

    logger.info(f"\n{'━'*60}")
    logger.info(f"  Crawling: {root}  (depth≤{MAX_DEPTH})")
    logger.info(f"{'━'*60}")

    # BFS queue: (url, depth)
    queue: deque = deque()
    queue.append((root, 0))

    pages_done = 0
    docs_done  = 0
    doc_queue: List[Tuple[str, str]] = []  # (doc_url, page_title)

    while queue and pages_done < MAX_PAGES_PER_SITE:
        url, depth = queue.popleft()
        norm = normalise_url(url)

        if norm in visited_urls:
            continue
        visited_urls.add(norm)

        if depth > MAX_DEPTH:
            continue

        time.sleep(DELAY_HTML)
        final_url, soup = fetch_html(url)
        if not soup:
            continue

        pages_done += 1
        logger.info(f"  [d{depth}] ({pages_done}) {url}")

        # Extract Q&A from this page
        page_pairs = extract_html_qa(soup, final_url or url)
        all_pairs.extend(page_pairs)
        logger.info(f"    → {len(page_pairs)} Q&A pairs")

        # Discover links
        html_links, doc_links = extract_links(soup, final_url or url, origin)

        # Queue new HTML pages
        for link in sorted(html_links):
            n = normalise_url(link)
            if n not in visited_urls:
                queue.append((link, depth + 1))

        # Collect doc links for later processing
        page_title = clean_text(soup.title.string) if (soup.title and soup.title.string) else site_label
        for dlink in doc_links:
            dn = normalise_url(dlink)
            if dn not in visited_docs:
                doc_queue.append((dlink, page_title))

    # ── Process documents ─────────────────────────────────────────────────
    logger.info(f"  Found {len(doc_queue)} documents to parse")

    for doc_url, page_title in doc_queue:
        if docs_done >= MAX_DOCS_PER_SITE:
            logger.info(f"  Doc limit ({MAX_DOCS_PER_SITE}) reached for {site_label}")
            break

        dn = normalise_url(doc_url)
        if dn in visited_docs:
            continue
        visited_docs.add(dn)

        ext = url_extension(doc_url)
        logger.info(f"  📄 Parsing {ext.upper()}: {doc_url}")
        time.sleep(DELAY_DOC)

        data = fetch_bytes(doc_url)
        if not data:
            continue

        text = parse_document(data, doc_url)
        if not text or len(text) < 80:
            logger.info(f"    → Empty or too short, skipping")
            continue

        doc_pairs = doc_text_to_qa(text, doc_url, page_title)
        logger.info(f"    → {len(doc_pairs)} Q&A chunks from document")
        all_pairs.extend(doc_pairs)
        docs_done += 1

    logger.info(f"  Site done: {pages_done} pages, {docs_done} docs, {len(all_pairs)} raw pairs")
    return all_pairs


# ─────────────────────────────────────────────────────────────────────────────
# DEDUPLICATION & QUALITY FILTER
# ─────────────────────────────────────────────────────────────────────────────

def dedup_and_filter(pairs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen: Set[str] = set()
    out  = []
    for p in pairs:
        q = (p.get("question") or "").strip()
        a = (p.get("answer")   or "").strip()

        if len(q) < 10 or len(a) < 20:
            continue

        key = text_hash(q)
        if key in seen:
            continue
        seen.add(key)

        # Clean up the pair
        p["question"] = q
        p["answer"]   = a
        out.append(p)

    return out


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def run_scraper(output_path: str = "data/gndec_data.json"):
    os.makedirs("data", exist_ok=True)

    all_pairs:    List[Dict[str, Any]] = []
    visited_urls: Set[str]             = set()
    visited_docs: Set[str]             = set()
    known_origins: Set[str]            = set(SEED_SITES)

    # ── Step 1: Discover extra subdomains from main page ──────────────────
    logger.info("🔍 Discovering subdomains from main GNDEC page...")
    _, main_soup = fetch_html(ROOT_URL)
    if main_soup:
        extra = discover_subdomains(main_soup, ROOT_URL, known_origins)
        logger.info(f"  Found {len(extra)} extra subdomains: {extra}")
        known_origins.update(extra)

    all_sites = list(dict.fromkeys(SEED_SITES + [
        s for s in known_origins if s not in set(SEED_SITES)
    ]))
    logger.info(f"Total sites to crawl: {len(all_sites)}")

    # ── Step 2: Crawl each site ───────────────────────────────────────────
    for idx, site in enumerate(all_sites, 1):
        logger.info(f"\n[{idx}/{len(all_sites)}] {site}")
        try:
            site_pairs = crawl_site(site, visited_urls, visited_docs)
            all_pairs.extend(site_pairs)
        except Exception as e:
            logger.error(f"Site {site} crashed: {e}", exc_info=True)

        # Incremental save after each site
        _save(all_pairs, output_path)
        logger.info(f"  💾 Running total: {len(all_pairs)} raw pairs")
        time.sleep(1.5)

    # ── Step 3: Dedup + filter ────────────────────────────────────────────
    before = len(all_pairs)
    all_pairs = dedup_and_filter(all_pairs)
    logger.info(f"\nDedup+filter: {before} → {len(all_pairs)} pairs")

    # ── Step 4: Final save ────────────────────────────────────────────────
    _save(all_pairs, output_path)
    logger.info(f"\n✅ Done! {len(all_pairs)} Q&A pairs saved to {output_path}")
    logger.info(f"   Pages crawled : {len(visited_urls)}")
    logger.info(f"   Docs parsed   : {len(visited_docs)}")

    return all_pairs


def _save(pairs: List[Dict[str, Any]], path: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(pairs, f, indent=2, ensure_ascii=False)


if __name__ == "__main__":
    run_scraper()
